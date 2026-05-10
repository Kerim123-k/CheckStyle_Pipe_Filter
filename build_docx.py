#!/usr/bin/env python3
"""Render report.md to .docx matching docs/ref/...EXAMPLE_WrongOne_But_Full_Right_format.docx style.

Style profile copied from the reference:
  body:    Arial 11pt
  H1:      Arial 20pt, black
  H2:      Arial 16pt
  H3:      Arial 14pt, #434343
  H4..H6:  Arial 12/11/11pt, #666666
  margins: 1 inch
  Code blocks rendered as single-cell shaded tables (matches the
  reference doc's 28 1x1 tables used as call-out boxes).
"""
from pathlib import Path

import markdown
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from htmldocx import HtmlToDocx

ROOT = Path(__file__).parent
SRC = ROOT / "report.md"
OUT = ROOT / "docs" / "SENG326-Task2-Report-Team13.docx"

BODY_FONT = "Arial"
H_COLORS = {1: None, 2: None, 3: "434343", 4: "666666", 5: "666666", 6: "666666"}
H_SIZES = {1: 20, 2: 16, 3: 14, 4: 12, 5: 11, 6: 11}


def set_run_font(run, name=BODY_FONT, size_pt=None, color_hex=None, bold=None):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if color_hex is not None:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    if bold is not None:
        run.font.bold = bold


def configure_styles(doc):
    for sec in doc.sections:
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1)
        sec.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), BODY_FONT)

    for level, size in H_SIZES.items():
        s = styles[f"Heading {level}"]
        s.font.name = BODY_FONT
        s.font.size = Pt(size)
        s.font.bold = level == 1
        if H_COLORS[level]:
            s.font.color.rgb = RGBColor.from_string(H_COLORS[level])

    code = styles.add_style("CodeBlock", 1) if "CodeBlock" not in [st.name for st in styles] else styles["CodeBlock"]
    code.font.name = "Roboto Mono"
    code.font.size = Pt(9)


def add_cover(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("SENG 326 — Software Architecture")
    set_run_font(r, size_pt=26, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Task 2 — Architectural Refactoring of Checkstyle 13.2.0")
    set_run_font(r, size_pt=15, color_hex="666666")

    doc.add_paragraph()
    doc.add_paragraph()

    target = doc.add_paragraph()
    target.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = target.add_run("Target architecture: Pipe-and-Filter")
    set_run_font(r, size_pt=14, bold=True)

    doc.add_paragraph()

    for line in [
        "KARIM HARIRI – 22050941008",
        "MOHAMED ATTIA EID ATTIA EID – 22050941017",
        "ZAID HARDAN – 22050941005",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        set_run_font(r, size_pt=12)

    doc.add_page_break()


def strip_cover_from_md(md_text):
    """Drop the front-matter we render manually so it isn't duplicated."""
    lines = md_text.splitlines()
    # find first '## Executive Summary' and start from there
    for i, line in enumerate(lines):
        if line.strip().lower().startswith("## executive summary"):
            return "\n".join(lines[i:])
    return md_text


md_text = strip_cover_from_md(SRC.read_text(encoding="utf-8"))
html = markdown.markdown(
    md_text,
    extensions=["fenced_code", "tables", "toc", "sane_lists"],
)

doc = Document()
configure_styles(doc)
add_cover(doc)

converter = HtmlToDocx()
converter.table_style = "Light Grid Accent 1"
converter.add_html_to_document(html, doc)

# Re-apply Arial to every run (htmldocx leaves font name unset)
for p in doc.paragraphs:
    for r in p.runs:
        if not r.font.name:
            set_run_font(r, name=BODY_FONT)
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    if not r.font.name:
                        set_run_font(r, name=BODY_FONT, size_pt=10)

doc.save(str(OUT))
print(f"wrote {OUT}")
